import os
import tempfile
import subprocess
import shutil
import time
import re
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
import whisper
from werkzeug.utils import secure_filename
import mimetypes

# Optional imports for document conversion and OCR
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader, PdfWriter
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook, Workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
    
    # Try to find Tesseract executable in common Windows locations
    import os as os_module
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(os_module.getenv('USERNAME', '')),
    ]
    
    # Check if tesseract is in PATH
    tesseract_found = False
    if shutil.which('tesseract'):
        tesseract_found = True
    else:
        # Try common installation paths
        for path in tesseract_paths:
            if os_module.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                tesseract_found = True
                print(f"Found Tesseract at: {path}")
                break
    
    if not tesseract_found:
        print("Warning: Tesseract OCR not found in PATH or common locations.")
        print("OCR functionality may not work. Please install Tesseract OCR.")
        print("Download from: https://github.com/UB-Mannheim/tesseract/wiki")
        
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_AUDIO_EXTENSIONS = {
    'mp3', 'wav', 'm4a', 'flac', 'ogg', 'opus', 'aac', 
    'wma', 'mp4', 'webm', '3gp', 'amr', 'aiff', 'au'
}
ALLOWED_DOCUMENT_EXTENSIONS = {
    'docx', 'doc', 'pdf', 'txt', 'rtf', 'xlsx', 'xls', 'csv', 'md', 'html'
}
ALLOWED_IMAGE_EXTENSIONS = {
    'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif', 'webp'
}
ALLOWED_EXTENSIONS = ALLOWED_AUDIO_EXTENSIONS | ALLOWED_DOCUMENT_EXTENSIONS | ALLOWED_IMAGE_EXTENSIONS
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs('transcriptions', exist_ok=True)
os.makedirs('conversions', exist_ok=True)
os.makedirs('ocr_results', exist_ok=True)

# Load Whisper model (base model for good balance of speed and accuracy)
# You can change this to 'tiny', 'small', 'medium', or 'large' based on your needs
model = None

def check_ffmpeg():
    """Check if FFmpeg is installed and available"""
    ffmpeg_path = shutil.which('ffmpeg')
    if ffmpeg_path is None:
        return False, None
    try:
        # Try to get version to verify it works
        result = subprocess.run(['ffmpeg', '-version'], 
                              capture_output=True, 
                              timeout=5,
                              text=True)
        if result.returncode == 0:
            return True, ffmpeg_path
    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        pass
    return False, None

def load_model():
    global model
    if model is None:
        print("Loading Whisper model... This may take a moment.")
        load_start = time.time()
        model = whisper.load_model("base")
        load_end = time.time()
        load_time = load_end - load_start
        print(f"Model loaded successfully! (took {load_time:.2f} seconds)")
        return model, load_time
    return model, 0

def format_transcription_with_sentences(text):
    """
    Format transcription text with sentences on separate lines.
    Splits on sentence endings (. ! ?) followed by space and capital letter.
    """
    if not text:
        return text
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Split on sentence endings: . ! ? followed by space and capital letter
    pattern = r'([.!?])\s+([A-Z])'
    parts = re.split(pattern, text)
    
    if len(parts) == 1:
        # No sentence endings found, return as is
        return text
    
    # Reconstruct sentences
    formatted_lines = []
    i = 0
    
    while i < len(parts):
        if i == 0:
            # First sentence (before any split)
            if parts[i].strip():
                formatted_lines.append(parts[i].strip())
        elif i + 1 < len(parts):
            # We have: punctuation, space, next sentence start
            punctuation = parts[i]  # . ! or ?
            next_sentence_start = parts[i + 1]  # Capital letter + rest
            
            # Complete previous sentence with punctuation
            if formatted_lines:
                formatted_lines[-1] += punctuation
            else:
                formatted_lines.append(punctuation)
            
            # Start new sentence
            if next_sentence_start.strip():
                formatted_lines.append(next_sentence_start.strip())
            
            i += 1  # Skip next part as we processed it
        else:
            # Remaining text
            if parts[i].strip():
                if formatted_lines:
                    formatted_lines[-1] += " " + parts[i].strip()
                else:
                    formatted_lines.append(parts[i].strip())
        
        i += 1
    
    # Join with newlines
    result = '\n'.join(line.strip() for line in formatted_lines if line.strip())
    
    # Clean up excessive newlines
    result = re.sub(r'\n{3,}', '\n\n', result)
    
    return result

def allowed_file(filename, extension_set=None):
    if extension_set is None:
        extension_set = ALLOWED_EXTENSIONS
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in extension_set

def get_file_extension(filename):
    """Get file extension without dot"""
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def convert_docx_to_txt(filepath):
    """Convert DOCX to TXT"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available")
    
    doc = Document(filepath)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

def convert_docx_to_pdf(filepath, output_path):
    """Convert DOCX to PDF"""
    if not DOCX_AVAILABLE or not REPORTLAB_AVAILABLE:
        raise Exception("Required libraries not available (python-docx or reportlab)")
    
    doc = Document(filepath)
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            p = Paragraph(paragraph.text, styles['Normal'])
            story.append(p)
    
    pdf.build(story)
    return output_path

def convert_pdf_to_txt(filepath):
    """Convert PDF to TXT"""
    if not PDF_AVAILABLE:
        raise Exception("pypdf library not available")
    
    reader = PdfReader(filepath)
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return '\n'.join(text)

def convert_pdf_to_docx(filepath, output_path):
    """Convert PDF to DOCX"""
    if not PDF2DOCX_AVAILABLE:
        raise Exception("pdf2docx library not available")
    
    cv = Converter(filepath)
    cv.convert(output_path)
    cv.close()
    return output_path

def convert_txt_to_pdf(filepath, output_path):
    """Convert TXT to PDF"""
    if not REPORTLAB_AVAILABLE:
        raise Exception("reportlab library not available")
    
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for line in text.split('\n'):
        if line.strip():
            p = Paragraph(line, styles['Normal'])
            story.append(p)
    
    pdf.build(story)
    return output_path

def convert_txt_to_docx(filepath, output_path):
    """Convert TXT to DOCX"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available")
    
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    
    doc = Document()
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    doc.save(output_path)
    return output_path

def convert_excel_to_txt(filepath):
    """Convert Excel to TXT"""
    if not EXCEL_AVAILABLE:
        raise Exception("openpyxl library not available")
    
    wb = load_workbook(filepath)
    text_lines = []
    
    for sheet_name in wb.sheetnames:
        text_lines.append(f"Sheet: {sheet_name}")
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
            if row_text.strip():
                text_lines.append(row_text)
        text_lines.append('')
    
    return '\n'.join(text_lines)

def perform_ocr(filepath):
    """Perform OCR on image or PDF file"""
    if not OCR_AVAILABLE:
        raise Exception("pytesseract library not available")
    
    ext = get_file_extension(filepath)
    text_parts = []
    
    if ext == 'pdf':
        if not PDF2IMAGE_AVAILABLE:
            raise Exception("pdf2image library not available for PDF OCR")
        
        # Convert PDF pages to images
        images = convert_from_path(filepath)
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image)
            if text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{text}")
    else:
        # Image file
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
        text_parts.append(text)
    
    return '\n\n'.join(text_parts)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'healthy'})

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({
            'error': f'File type not allowed. Supported formats: {", ".join(ALLOWED_EXTENSIONS)}'
        }), 400
    
    # Check for FFmpeg before processing
    ffmpeg_available, ffmpeg_path = check_ffmpeg()
    if not ffmpeg_available:
        return jsonify({
            'error': 'FFmpeg is not installed or not found in PATH. FFmpeg is required for audio transcription. Please install FFmpeg and restart the server. See INSTALL_FFMPEG.md for installation instructions.'
        }), 500
    
    try:
        # Start timing
        start_time = time.time()
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Check file size
        file_size = os.path.getsize(filepath)
        if file_size > MAX_FILE_SIZE:
            os.remove(filepath)
            return jsonify({'error': f'File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024)}MB'}), 400
        
        # Load model and transcribe
        whisper_model, model_load_time = load_model()
        print(f"Transcribing {filename}...")
        
        # Transcribe audio
        transcription_start = time.time()
        result = whisper_model.transcribe(
            filepath,
            language=None,  # Auto-detect language
            task="transcribe"
        )
        transcription_end = time.time()
        
        # Extract transcription text
        transcription_text = result["text"]
        
        # Format transcription with sentences on separate lines
        formatted_text = format_transcription_with_sentences(transcription_text)
        
        # Save transcription to file
        transcription_filename = os.path.splitext(filename)[0] + '.txt'
        transcription_path = os.path.join('transcriptions', transcription_filename)
        
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(formatted_text)
        
        # Clean up uploaded file
        os.remove(filepath)
        
        # Calculate processing time
        end_time = time.time()
        processing_time = end_time - start_time
        transcription_time = transcription_end - transcription_start
        
        return jsonify({
            'success': True,
            'transcription': formatted_text,
            'filename': transcription_filename,
            'language': result.get('language', 'unknown'),
            'download_url': f'/download/{transcription_filename}',
            'processing_time': round(processing_time, 2),
            'transcription_time': round(transcription_time, 2),
            'model_load_time': round(model_load_time, 2)
        })
    
    except Exception as e:
        # Clean up on error
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'Transcription failed: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        filepath = os.path.join('transcriptions', secure_filename(filename))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/supported-formats', methods=['GET'])
def supported_formats():
    ffmpeg_available, ffmpeg_path = check_ffmpeg()
    return jsonify({
        'formats': list(ALLOWED_EXTENSIONS),
        'max_file_size_mb': MAX_FILE_SIZE / (1024 * 1024),
        'ffmpeg_available': ffmpeg_available,
        'ffmpeg_path': ffmpeg_path
    })

@app.route('/check-ffmpeg', methods=['GET'])
def check_ffmpeg_status():
    """Check FFmpeg installation status"""
    ffmpeg_available, ffmpeg_path = check_ffmpeg()
    if ffmpeg_available:
        return jsonify({
            'available': True,
            'path': ffmpeg_path,
            'message': 'FFmpeg is installed and ready'
        })
    else:
        return jsonify({
            'available': False,
            'path': None,
            'message': 'FFmpeg is not installed or not found in PATH. Please install FFmpeg to use transcription features.',
            'install_guide': 'See INSTALL_FFMPEG.md for installation instructions'
        }), 503

@app.route('/convert-document', methods=['POST'])
def convert_document():
    """Convert document from one format to another"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    target_format = request.form.get('target_format', '').lower()
    if not target_format:
        return jsonify({'error': 'Target format not specified'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    try:
        start_time = time.time()
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        file_size = os.path.getsize(filepath)
        if file_size > MAX_FILE_SIZE:
            os.remove(filepath)
            return jsonify({'error': f'File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024)}MB'}), 400
        
        source_ext = get_file_extension(filename)
        
        # Generate output filename
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}.{target_format}"
        output_path = os.path.join('conversions', output_filename)
        
        # Perform conversion based on source and target formats
        if source_ext == 'docx' and target_format == 'txt':
            text = convert_docx_to_txt(filepath)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif source_ext == 'docx' and target_format == 'pdf':
            convert_docx_to_pdf(filepath, output_path)
        elif source_ext == 'pdf' and target_format == 'txt':
            text = convert_pdf_to_txt(filepath)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif source_ext == 'pdf' and target_format == 'docx':
            convert_pdf_to_docx(filepath, output_path)
        elif source_ext == 'txt' and target_format == 'pdf':
            convert_txt_to_pdf(filepath, output_path)
        elif source_ext == 'txt' and target_format == 'docx':
            convert_txt_to_docx(filepath, output_path)
        elif source_ext in ['xlsx', 'xls'] and target_format == 'txt':
            text = convert_excel_to_txt(filepath)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        else:
            os.remove(filepath)
            return jsonify({'error': f'Conversion from {source_ext} to {target_format} is not supported'}), 400
        
        processing_time = time.time() - start_time
        
        # Clean up uploaded file
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'download_url': f'/download-conversion/{output_filename}',
            'processing_time': round(processing_time, 2)
        })
    
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

@app.route('/ocr', methods=['POST'])
def ocr_endpoint():
    """Perform OCR on image or PDF file"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not OCR_AVAILABLE:
        return jsonify({
            'error': 'OCR functionality not available. Please install pytesseract and Tesseract OCR engine.'
        }), 503
    
    try:
        start_time = time.time()
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        file_size = os.path.getsize(filepath)
        if file_size > MAX_FILE_SIZE:
            os.remove(filepath)
            return jsonify({'error': f'File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024)}MB'}), 400
        
        # Perform OCR
        ocr_text = perform_ocr(filepath)
        
        # Format with sentences
        formatted_text = format_transcription_with_sentences(ocr_text)
        
        # Save OCR result
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}_ocr.txt"
        output_path = os.path.join('ocr_results', output_filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_text)
        
        processing_time = time.time() - start_time
        
        # Clean up uploaded file
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'text': formatted_text,
            'filename': output_filename,
            'download_url': f'/download-ocr/{output_filename}',
            'processing_time': round(processing_time, 2)
        })
    
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': f'OCR failed: {str(e)}'}), 500

@app.route('/download-conversion/<filename>')
def download_conversion(filename):
    """Download converted document"""
    try:
        filepath = os.path.join('conversions', secure_filename(filename))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download-ocr/<filename>')
def download_ocr(filename):
    """Download OCR result"""
    try:
        filepath = os.path.join('ocr_results', secure_filename(filename))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/supported-conversions', methods=['GET'])
def supported_conversions():
    """Get supported document conversion formats"""
    return jsonify({
        'conversions': {
            'docx': ['txt', 'pdf'],
            'pdf': ['txt', 'docx'],
            'txt': ['pdf', 'docx'],
            'xlsx': ['txt'],
            'xls': ['txt']
        },
        'libraries_available': {
            'docx': DOCX_AVAILABLE,
            'pdf': PDF_AVAILABLE,
            'pdf2docx': PDF2DOCX_AVAILABLE,
            'excel': EXCEL_AVAILABLE,
            'reportlab': REPORTLAB_AVAILABLE
        }
    })

@app.route('/ocr-capabilities', methods=['GET'])
def ocr_capabilities():
    """Get OCR capabilities and status"""
    return jsonify({
        'available': OCR_AVAILABLE,
        'pdf2image_available': PDF2IMAGE_AVAILABLE,
        'supported_formats': list(ALLOWED_IMAGE_EXTENSIONS) + ['pdf'] if OCR_AVAILABLE else [],
        'message': 'OCR is available' if OCR_AVAILABLE else 'OCR requires pytesseract and Tesseract OCR engine installation'
    })

if __name__ == '__main__':
    print("Starting Audio Transcription Server...")
    print("Supported formats:", ", ".join(ALLOWED_EXTENSIONS))
    
    # Check for FFmpeg at startup
    ffmpeg_available, ffmpeg_path = check_ffmpeg()
    if ffmpeg_available:
        print(f"✓ FFmpeg found at: {ffmpeg_path}")
    else:
        print("⚠ WARNING: FFmpeg not found!")
        print("  FFmpeg is required for audio transcription.")
        print("  Please install FFmpeg and restart the server.")
        print("  See INSTALL_FFMPEG.md for installation instructions.")
        print("  The server will start, but transcription will fail until FFmpeg is installed.")
    
    app.run(debug=True, host='0.0.0.0', port=5012)

